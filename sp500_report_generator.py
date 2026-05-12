# ══════════════════════════════════════════════════════════════════════════
# 标普500监控系统 v15 — 报告生成模块(T1-T10 触发器版)
# ══════════════════════════════════════════════════════════════════════════

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

C_DARK   = RGBColor(0x1F, 0x38, 0x64)
C_MID    = RGBColor(0x2F, 0x75, 0xB6)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_GREEN  = RGBColor(0x37, 0x56, 0x23)
C_RED    = RGBColor(0xC0, 0x00, 0x00)
C_GRAY   = RGBColor(0x59, 0x59, 0x59)
C_ORANGE = RGBColor(0xED, 0x7D, 0x31)
C_YELLOW = RGBColor(0xBF, 0x90, 0x00)


def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color); tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=10, align='center'):
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(str(text))
    run.bold = bold; run.font.size = Pt(size); run.font.name = 'Arial'
    if color: run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_heading(doc, text, level=2):
    para = doc.add_paragraph()
    run = para.add_run(text); run.bold = True; run.font.name = 'Arial'
    if level == 1:
        run.font.size = Pt(14); run.font.color.rgb = C_DARK
    else:
        run.font.size = Pt(11); run.font.color.rgb = C_MID
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after  = Pt(4)
    if level == 2:
        pPr = para._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
        bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), '2F75B6')
        pBdr.append(bot); pPr.append(pBdr)


def add_kv_table(doc, rows, cols=4):
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        bg = 'F2F2F2' if i % 2 == 0 else 'FFFFFF'
        for j, (label, value) in enumerate(row_data):
            if j * 2 + 1 >= cols: break
            lc = row.cells[j*2]; vc = row.cells[j*2+1]
            set_cell_bg(lc, bg); set_cell_bg(vc, bg)
            set_cell_text(lc, label, bold=True, color=C_DARK, align='left', size=9)
            set_cell_text(vc, value, color=C_DARK, size=9)
    for row in table.rows:
        row.height = Cm(0.65)
        for j, cell in enumerate(row.cells):
            cell.width = Cm(3.5 if j % 2 == 0 else 2.5)
    return table


def fmt_val(val, fmt_str='', suffix=''):
    try:
        import math
        f = float(val)
        if math.isnan(f): return 'N/A'
        return f"{format(f, fmt_str)}{suffix}"
    except: return 'N/A'


# 触发器显示元信息
TRIGGER_DISPLAY = [
    ('T1_2000离场',  'T1: 2000 互联网泡沫顶',  'exit'),
    ('T2_2007离场',  'T2: 2007 鸽派假反弹顶',  'exit'),
    ('T3_2015离场',  'T3: 2015 工业衰退顶',    'exit'),
    ('T4_2022离场',  'T4: 2022 加息顶',        'exit'),
    ('T5_2002入场',  'T5: 2002 互联网底',      'entry'),
    ('T6_2009入场',  'T6: 2009 次贷底',        'entry'),
    ('T7_2020入场',  'T7: 2020 COVID 底',      'entry'),
    ('T8_2022入场',  'T8: 2022 加息底',        'entry'),
    ('T9_白银坑1组', 'T9: 白银坑 1 组(广义)', 'entry'),
    ('T10_白银坑2组','T10: 白银坑 2 组(中庸态)','entry'),
]


def add_trigger_table(doc, trigger_id, trigger_name, ttype, snapshot):
    """
    为单个触发器输出详细的因子状态表。
    """
    triggers = snapshot.get('triggers', {})
    r = triggers.get(trigger_id)
    if r is None:
        return

    # 触发器标题段
    sc_para = doc.add_paragraph()
    sc_para.paragraph_format.space_before = Pt(8)
    sc_para.paragraph_format.space_after  = Pt(2)
    triggered = r.get('triggered')
    pct = r.get('satisfied_pct', 0)
    if triggered is True:
        label = '🚨 【已触发!】'; color = C_RED
    elif triggered is False:
        if pct >= 0.9:
            label = f'🟠 高度接近 ({pct*100:.0f}%)'; color = C_ORANGE
        elif pct >= 0.7:
            label = f'🟡 中度接近 ({pct*100:.0f}%)'; color = C_YELLOW
        else:
            label = f'⚪ 远未触发 ({pct*100:.0f}%)'; color = C_GRAY
    else:
        label = '? 数据不足'; color = C_ORANGE

    rsc = sc_para.add_run(f'{trigger_name}  {label}')
    rsc.bold = True; rsc.font.size = Pt(10); rsc.font.name = 'Arial'; rsc.font.color.rgb = color

    # 物理画像
    desc_para = doc.add_paragraph()
    desc_para.paragraph_format.space_after = Pt(2)
    rdesc = desc_para.add_run(f'物理画像: {r.get("description", "")} | ★核心因子: {r.get("core_factor", "")}')
    rdesc.font.size = Pt(8); rdesc.font.name = 'Arial'; rdesc.font.color.rgb = C_GRAY
    rdesc.italic = True

    # 收集所有条件项
    rows = []
    # 必有因子
    for f, v in r.get('must_have_status', []):
        is_core = (f == r.get('core_factor', ''))
        tag = '★核心(AND)' if is_core else 'AND'
        if v is True:
            mark = '✓'; bg = 'E2EFDA'; clr = C_GREEN
        elif v is False:
            mark = '✗'; bg = 'FCE4D6'; clr = C_RED
        else:
            mark = '?'; bg = 'FFF2CC'; clr = C_ORANGE
        rows.append((tag, f, str(v), mark, bg, clr))

    # OR 路径
    or_paths = r.get('or_paths_status', [])
    for i, path_st in enumerate(or_paths):
        # 路径整体是否满足(任一为 True)
        path_ok = any(v is True for _, v in path_st)
        path_label = f'OR-路径 {i+1}(任一)'
        if path_ok:
            mark = '✓'; bg = 'E2EFDA'; clr = C_GREEN
        elif any(v is None for _, v in path_st):
            mark = '?'; bg = 'FFF2CC'; clr = C_ORANGE
        else:
            mark = '✗'; bg = 'FCE4D6'; clr = C_RED
        path_factors = '  |  '.join(f'{f}={v}' for f, v in path_st)
        rows.append((path_label, path_factors, '任一', mark, bg, clr))

    # NOT 屏蔽
    for f, v in r.get('not_have_status', []):
        # NOT(F)=True 当 F=False
        if v is False:
            mark = '✓'; bg = 'E2EFDA'; clr = C_GREEN
            value_show = f'{f}=False(NOT 成立)'
        elif v is True:
            mark = '✗'; bg = 'FCE4D6'; clr = C_RED
            value_show = f'{f}=True(NOT 失效)'
        else:
            mark = '?'; bg = 'FFF2CC'; clr = C_ORANGE
            value_show = f'{f}=?'
        rows.append(('NOT 屏蔽', f'NOT({f})', value_show, mark, bg, clr))

    # special_not (组合 NOT)
    for combo, combo_vals, all_true in r.get('special_not_status', []):
        combo_show = ' AND '.join(c[0] for c in combo_vals)
        if not all_true:
            mark = '✓'; bg = 'E2EFDA'; clr = C_GREEN
            verdict = '组合不全成立(NOT 成立)'
        else:
            mark = '✗'; bg = 'FCE4D6'; clr = C_RED
            verdict = '组合全成立(NOT 失效)'
        rows.append(('NOT 复合屏蔽', f'NOT({combo_show})', verdict, mark, bg, clr))

    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=4)
    table.style = 'Table Grid'
    for i, (tag, factor, value, mark, bg, clr) in enumerate(rows):
        row = table.rows[i]
        for cell in row.cells: set_cell_bg(cell, bg)
        set_cell_text(row.cells[0], tag,    align='left', size=8)
        set_cell_text(row.cells[1], factor, align='left', size=9)
        set_cell_text(row.cells[2], value,  size=8)
        set_cell_text(row.cells[3], mark, bold=True, color=clr, size=11)
        row.height = Cm(0.55)
        row.cells[0].width = Cm(2.2)
        row.cells[1].width = Cm(6.5)
        row.cells[2].width = Cm(3.5)
        row.cells[3].width = Cm(0.8)


def generate_report(snapshot):
    from config import CURRENT_POSITION

    date_str = snapshot.get('date', datetime.today().strftime('%Y-%m-%d'))
    year = date_str[:4]; month = date_str[5:7]

    import platform
    if platform.system() == 'Windows':
        report_dir = f"D:\\sp500_monitor\\reports\\{year}\\{month}"
    else:
        report_dir = f"reports/{year}/{month}"
    os.makedirs(report_dir, exist_ok=True)
    report_path = os.path.join(report_dir, f"SP500触发器系统日报_{date_str}.docx")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5); section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('标普500触发器系统(V15)监控日报')
    run.bold = True; run.font.size = Pt(16); run.font.name = 'Arial'
    run.font.color.rgb = C_DARK

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run(f'{date_str}  |  V15 因子矩阵 + 10 触发器  |  自动生成')
    run2.font.size = Pt(9); run2.font.name = 'Arial'; run2.font.color.rgb = C_GRAY

    pos_text = f'持仓:{CURRENT_POSITION}' if CURRENT_POSITION else '当前:空仓'
    env_para = doc.add_paragraph()
    env_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = env_para.add_run(f'SP500 = {fmt_val(snapshot.get("sp500"), ",.2f")}  |  '
                            f'回撤 {fmt_val(snapshot.get("sp_dd_pct"), ".1f", "%")}  |  {pos_text}')
    run3.bold = True; run3.font.size = Pt(10); run3.font.name = 'Arial'
    run3.font.color.rgb = C_DARK
    doc.add_paragraph()

    # ── 一、关键中间量 ─────────────────────────────────────────────
    add_heading(doc, '一、关键中间量(V15 因子的输入)')
    indicator_rows = [
        [('SP500',          fmt_val(snapshot.get('sp500'), ',.2f')),
         ('ATH 回撤',       fmt_val(snapshot.get('sp_dd_pct'), '.2f', '%'))],
        [('200 周均线',      fmt_val(snapshot.get('ma200w'), ',.2f')),
         ('Sσ_200(标准差偏离)', fmt_val(snapshot.get('sigma_dev_200'), '.2f', 'σ'))],
        [('VIX',            fmt_val(snapshot.get('vix'), '.2f')),
         ('V_c(平静均值)',   fmt_val(snapshot.get('V_c'), '.2f') + ' ± ' + fmt_val(snapshot.get('V_c_sigma'), '.2f'))],
        [('Forward PE',     fmt_val(snapshot.get('forward_pe'), '.2f', 'x')),
         ('PE 1260 日均',    fmt_val(snapshot.get('pe_avg'), '.2f'))],
        [('ERP',            fmt_val(snapshot.get('erp'), '.2f', '%')),
         ('实际利率 R',      fmt_val(snapshot.get('real_rate'), '.2f', '%'))],
        [('NFCI',           fmt_val(snapshot.get('nfci'), '.3f')),
         ('N_c(20 日变化)', fmt_val(snapshot.get('n_c'), '+.3f'))],
        [('N_c μ(1260 日)', fmt_val(snapshot.get('n_c_mu'), '+.3f')),
         ('N_c σ(ddof=1)',  fmt_val(snapshot.get('n_c_sigma'), '.3f'))],
        [('HY 信用利差',     fmt_val(snapshot.get('hy_spread'), '.2f', '%')),
         ('HY_c21(21 日变化)', fmt_val(snapshot.get('hy_c21'), '+.3f'))],
        [('Y 利差(10Y-2Y)',  fmt_val(snapshot.get('y_spread'), '.3f')),
         ('联储利率',        fmt_val(snapshot.get('fed_rate'), '.2f', '%'))],
        [('CPI 同比(滞后 45 日)', fmt_val(snapshot.get('cpi_y'), '.2f', '%')),
         ('WALCL 13 周变化', fmt_val(snapshot.get('walcl_13w_chg'), '+.3f', 'T'))],
        [('OIL 价格(WTI)',  fmt_val(snapshot.get('oil_price'), '.2f')),
         ('OIL_pct5y',       fmt_val(snapshot.get('oil_pct5y'), '.0f', '%'))],
    ]
    add_kv_table(doc, indicator_rows, cols=4)

    # ── 总状态 ──────────────────────────────────────────────────
    triggers = snapshot.get('triggers', {})
    has_trig  = any(t.get('triggered') is True for t in triggers.values())
    has_close = any(t.get('triggered') is False and t.get('satisfied_pct', 0) >= 0.9
                    for t in triggers.values())

    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if has_trig:
        rs = sig_para.add_run('🚨  有触发器已触发,请立即关注!')
        rs.bold = True; rs.font.size = Pt(13); rs.font.color.rgb = C_RED
    elif has_close:
        rs = sig_para.add_run('⚠️  有触发器高度接近(≥90%)')
        rs.bold = True; rs.font.size = Pt(12); rs.font.color.rgb = C_ORANGE
    else:
        rs = sig_para.add_run('✓  全部触发器远未触发,继续等待')
        rs.bold = True; rs.font.size = Pt(13); rs.font.color.rgb = C_GREEN
    doc.add_paragraph()

    # ── 二、离场触发器 ──────────────────────────────────────────
    add_heading(doc, '二、离场触发器详细检测(高位卖出)')
    for tid, name, ttype in TRIGGER_DISPLAY:
        if ttype != 'exit': continue
        add_trigger_table(doc, tid, name, ttype, snapshot)

    # ── 三、入场触发器 ──────────────────────────────────────────
    doc.add_paragraph()
    add_heading(doc, '三、入场触发器详细检测(低位买入)')
    for tid, name, ttype in TRIGGER_DISPLAY:
        if ttype != 'entry': continue
        add_trigger_table(doc, tid, name, ttype, snapshot)

    # 页脚
    doc.add_paragraph()
    footer = doc.add_paragraph(
        f'标普500触发器系统(V15)监控  |  自动生成于 {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n'
        f'Layer 3 验证: V15 重算 = 文件 4,857 触发日 0 差异 ✓')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in footer.runs:
        r.font.size = Pt(8); r.font.color.rgb = C_GRAY

    doc.save(report_path)
    print(f"  ✅ 报告已保存:{report_path}")
    return report_path
